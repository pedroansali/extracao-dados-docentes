from docx import Document
import xlsxwriter

def pegar_dados_word(caminho_arquivo):

    #abrindo arquivo documento
    documento = open(caminho_arquivo, 'rb')

    #lendo o documento
    leitor_documento = Document(documento)
    data = ""

    for p in leitor_documento.paragraphs:
        data += p.text+"\n"

    return data

def componente(arquivo): #função para detectar a qual matéria se refere o relatório
    if arquivo.find('Matemática') != -1 or arquivo.find('Matematica') != -1 or arquivo.find('matemática') != -1 or arquivo.find('matematica') != -1:
        #relatório é de matemática
        componente = "Matemática"
    elif arquivo.find('Português') != -1 or arquivo.find('português') != -1 or arquivo.find('Portugues') != -1 or arquivo.find('portugues') != -1:
        #relatório é de português
        componente = "Português"
    return componente

def turma(arquivo): #função para detectar a qual turma se refere o relatório
    if arquivo.find('6º') != -1:
        turma = "6º Ano"
    elif arquivo.find('7º') != -1:
        turma = "7º Ano"
    elif arquivo.find('8º') != -1:
        turma = "8º Ano"
    elif arquivo.find('9º') != -1:
        turma = "9º Ano"
    elif arquivo.find('5º') != -1:
        turma = "5º Ano"
    elif arquivo.find('4º') != -1:
        turma = "4º Ano"
    elif arquivo.find('3º') != -1:
        turma = "3º Ano"
    elif arquivo.find('2º') != -1:
        turma = "2º Ano"
    elif arquivo.find('1º') != -1:
        turma = "1º Ano"
    else:
        turma = " "
    return turma

def encontrar_estrategia(caminho_arquivo):
    document = Document(caminho_arquivo)
    metodologia = ""
    # Find the index of the `Summary of Broadspectrum Offer` syntax and store it
    ind = [i for i, para in enumerate(document.paragraphs) if 'Estratégia da Aula:' in para.text]
    # Print the text for any element with an index greater than the index found in the list comprehension above
    if ind:
        for i, para in enumerate(document.paragraphs):
            if i > ind[0]:
                metodologia = metodologia + " " + para.text
    metodologia = metodologia.split(" Atividades")[0]
    return metodologia

def encontrar_recurso(caminho_arquivo):
    document = Document(caminho_arquivo)
    recurso_didatico = ""
    # Find the index of the `Summary of Broadspectrum Offer` syntax and store it
    ind = [i for i, para in enumerate(document.paragraphs) if 'Recursos da Aula:' in para.text]
    # Print the text for any element with an index greater than the index found in the list comprehension above
    if ind:
        for i, para in enumerate(document.paragraphs):
            if i > ind[0]:
                recurso_didatico = recurso_didatico + " " + para.text
    recurso_didatico = recurso_didatico.split(" Registro")[0]
    return recurso_didatico

caminho_arquivo = "teste/1 (" + "1" + ").docx"
arquivo = pegar_dados_word(caminho_arquivo)
print(arquivo)

workbook = xlsxwriter.Workbook('Teste.xlsx')
worksheet = workbook.add_worksheet()
row = 0
col = 0
worksheet.write(row, col, "TURMA")
worksheet.write(row, col + 1, "COMPONENTE")
worksheet.write(row, col + 2, "ESTRATÉGIA METODOLÓGICA")
worksheet.write(row, col + 3, "RECURSOS DIDÁTICOS")
for i in range(1,53): #PREENCHER COM NUMERO DE ARQUIVOS NA PASTA + 1
    caminho_arquivo = "teste/1 (" + str(i) + ").docx"
    arquivo = pegar_dados_word(caminho_arquivo)
    #print(arquivo)
    #componete = componente(arquivo)
    componete = "Matemática"
    turm = turma(arquivo)
    metodologia = encontrar_estrategia(caminho_arquivo)
    #print(metodologia)
    recurso_didatico = encontrar_recurso(caminho_arquivo)
    #print(recurso_didatico)

    row = i
    col = 0
    worksheet.write(row, col, turm)
    worksheet.write(row, col+1, componete)
    worksheet.write(row, col+2, metodologia)
    worksheet.write(row, col+3, recurso_didatico)

workbook.close()


